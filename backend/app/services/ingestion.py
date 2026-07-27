import io
import re
from pathlib import Path
from uuid import uuid4

import fitz
from docx import Document as DocxDocument
from qdrant_client import models
from sqlmodel import Session

from app.models import Document, KnowledgeBase
from app.schemas import IngestResponse
from app.services.rag import RagService

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class IngestionError(ValueError):
    """Raised when a submitted personal document cannot be indexed."""


def extract_text(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        accepted = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise IngestionError(f"Format non pris en charge. Formats acceptes : {accepted}")
    if extension == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as pdf:
            return "\n".join(page.get_text() for page in pdf)
    if extension == ".docx":
        document = DocxDocument(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise IngestionError("Le fichier texte doit etre encode en UTF-8.") from error


def split_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> list[str]:
    normalized_text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not normalized_text:
        return []

    chunks: list[str] = []
    start = 0
    text_length = len(normalized_text)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        if end < text_length:
            paragraph_boundary = normalized_text.rfind("\n\n", start + chunk_size // 2, end)
            sentence_boundary = normalized_text.rfind(". ", start + chunk_size // 2, end)
            boundary = max(paragraph_boundary, sentence_boundary)
            if boundary > start:
                end = boundary + (2 if boundary == paragraph_boundary else 1)
        chunk = normalized_text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_length:
            break
        start = max(end - overlap, start + 1)
    return chunks


class IngestionService:
    def __init__(self, rag_service: RagService):
        self.rag_service = rag_service

    async def ingest(
        self,
        session: Session,
        kb_id: int,
        filename: str,
        raw_text: str,
    ) -> IngestResponse:
        if not session.get(KnowledgeBase, kb_id):
            raise IngestionError("Base de connaissances introuvable.")

        chunks = split_text(raw_text)
        if not chunks:
            raise IngestionError("Aucun texte exploitable n'a ete trouve.")

        document = Document(kb_id=kb_id, filename=filename, raw_text=raw_text.strip())
        session.add(document)
        session.commit()
        session.refresh(document)

        try:
            embeddings = [await self.rag_service.ollama.embed(chunk) for chunk in chunks]
            await self.rag_service.ensure_collection(len(embeddings[0]))
            points = [
                models.PointStruct(
                    id=str(uuid4()),
                    vector=embedding,
                    payload={
                        "kb_id": kb_id,
                        "document_id": document.id,
                        "chunk_index": chunk_index,
                        "text": chunk,
                        "source": filename,
                    },
                )
                for chunk_index, (chunk, embedding) in enumerate(zip(chunks, embeddings, strict=True))
            ]
            await self.rag_service.qdrant.upsert(
                collection_name=self.rag_service.settings.qdrant_collection,
                points=points,
                wait=True,
            )
        except Exception:
            session.delete(document)
            session.commit()
            raise

        return IngestResponse(
            document_id=document.id,
            kb_id=kb_id,
            filename=filename,
            chunks_indexed=len(chunks),
        )